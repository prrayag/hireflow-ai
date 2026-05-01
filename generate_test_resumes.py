"""
generate_test_resumes.py
Generates 10 realistic test resumes in mixed formats (PDF + DOCX)
and zips them into test_data/test_resumes.zip
"""

import os, zipfile
from pathlib import Path

# ── Output folder ────────────────────────────────────────────────────────────
OUT_DIR = Path(__file__).parent / "test_data"
OUT_DIR.mkdir(exist_ok=True)

# ── 10 Candidate profiles ────────────────────────────────────────────────────
CANDIDATES = [
    {
        "name": "Arjun Mehta",
        "email": "arjun.mehta@gmail.com",
        "phone": "+91 98201 34567",
        "location": "Mumbai, Maharashtra",
        "role": "Senior Software Engineer",
        "experience": "6 years",
        "education": "B.Tech Computer Science — IIT Bombay, 2018",
        "skills": ["Python", "Django", "FastAPI", "PostgreSQL", "Redis", "Docker",
                   "Kubernetes", "AWS EC2", "S3", "React", "REST APIs", "Git"],
        "summary": ("Results-driven Senior Software Engineer with 6 years of experience "
                    "building scalable backend systems. Proficient in Python ecosystems, "
                    "cloud infrastructure, and leading cross-functional teams of 5+ engineers."),
        "jobs": [
            ("Razorpay", "Senior Software Engineer", "Jan 2022 – Present",
             ["Designed and maintained payment gateway APIs handling 2M+ daily transactions",
              "Reduced API latency by 38% through Redis caching and query optimisation",
              "Led migration from monolith to microservices architecture using Kubernetes",
              "Mentored 4 junior engineers and conducted bi-weekly code reviews"]),
            ("Infosys", "Software Engineer", "Jul 2018 – Dec 2021",
             ["Built RESTful APIs for an e-commerce platform serving 500K monthly users",
              "Implemented CI/CD pipelines using GitHub Actions and Docker",
              "Automated regression testing, cutting QA cycle time by 40%"]),
        ],
        "certs": ["AWS Certified Solutions Architect – Associate",
                  "Certified Kubernetes Administrator (CKA)"],
        "format": "pdf",
    },
    {
        "name": "Priya Sharma",
        "email": "priya.sharma@outlook.com",
        "phone": "+91 77003 89012",
        "location": "Bengaluru, Karnataka",
        "role": "Machine Learning Engineer",
        "experience": "4 years",
        "education": "M.Tech Artificial Intelligence — BITS Pilani, 2020",
        "skills": ["Python", "TensorFlow", "PyTorch", "scikit-learn", "Pandas",
                   "NumPy", "SQL", "Spark", "Airflow", "MLflow", "Hugging Face", "Docker"],
        "summary": ("ML Engineer passionate about building production-grade AI systems. "
                    "Experienced in end-to-end model development, MLOps, and NLP solutions "
                    "across fintech and healthtech domains."),
        "jobs": [
            ("PhonePe", "Machine Learning Engineer", "Mar 2021 – Present",
             ["Built fraud detection model reducing false positives by 22% (XGBoost + SHAP)",
              "Deployed NLP-based customer intent classifier with 91% accuracy in production",
              "Built Airflow DAGs for automated model retraining pipelines",
              "Integrated MLflow for experiment tracking across 15 ML projects"]),
            ("Mu Sigma", "Data Scientist", "Jul 2020 – Feb 2021",
             ["Performed exploratory analysis on 10M+ row datasets for retail client",
              "Built churn prediction models using logistic regression and random forest"]),
        ],
        "certs": ["Google Professional ML Engineer",
                  "Deep Learning Specialization – deeplearning.ai"],
        "format": "docx",
    },
    {
        "name": "Rohan Verma",
        "email": "rohan.verma@yahoo.com",
        "phone": "+91 90011 22334",
        "location": "Pune, Maharashtra",
        "role": "Full Stack Developer",
        "experience": "3 years",
        "education": "B.E. Information Technology — VIT Pune, 2021",
        "skills": ["React", "Next.js", "Node.js", "Express", "MongoDB", "PostgreSQL",
                   "TypeScript", "Tailwind CSS", "GraphQL", "Docker", "AWS Lambda", "Git"],
        "summary": ("Full Stack Developer specialising in React and Node.js ecosystems. "
                    "Strong focus on performance, clean UI/UX, and scalable API design. "
                    "Experience shipping SaaS products from zero to production."),
        "jobs": [
            ("Zoho Corporation", "Full Stack Developer", "Aug 2021 – Present",
             ["Built and maintained 3 internal SaaS modules used by 8,000+ employees",
              "Migrated legacy jQuery app to React 18 with TypeScript, improving load time by 55%",
              "Designed GraphQL API layer reducing over-fetching by 60%",
              "Implemented OAuth2 SSO integration across 5 Zoho products"]),
        ],
        "certs": ["Meta Frontend Developer Certificate",
                  "MongoDB Developer Associate"],
        "format": "pdf",
    },
    {
        "name": "Sneha Pillai",
        "email": "sneha.pillai@gmail.com",
        "phone": "+91 88109 45678",
        "location": "Hyderabad, Telangana",
        "role": "Data Analyst",
        "experience": "2 years",
        "education": "B.Sc Statistics — University of Hyderabad, 2022",
        "skills": ["Python", "Pandas", "SQL", "Tableau", "Power BI", "Excel",
                   "R", "NumPy", "Matplotlib", "Google Analytics", "Looker", "Snowflake"],
        "summary": ("Detail-oriented Data Analyst with a strong foundation in statistical "
                    "analysis and business intelligence. Experienced in translating complex "
                    "datasets into executive-level insights and dashboards."),
        "jobs": [
            ("Deloitte India", "Data Analyst", "Jul 2022 – Present",
             ["Created Power BI dashboards for 3 Fortune 500 clients tracking 50+ KPIs",
              "Wrote optimised SQL queries reducing report generation time from 2 hrs to 8 min",
              "Performed A/B test analysis that increased client email CTR by 18%",
              "Automated weekly reporting pipeline using Python + Google Sheets API"]),
        ],
        "certs": ["Google Data Analytics Certificate",
                  "Tableau Desktop Specialist"],
        "format": "docx",
    },
    {
        "name": "Karan Bhatia",
        "email": "karan.bhatia@protonmail.com",
        "phone": "+91 96300 77891",
        "location": "Delhi, NCR",
        "role": "DevOps Engineer",
        "experience": "5 years",
        "education": "B.Tech Electronics — DTU Delhi, 2019",
        "skills": ["Kubernetes", "Docker", "Terraform", "Ansible", "Jenkins", "GitHub Actions",
                   "AWS", "GCP", "Azure", "Linux", "Bash", "Python", "Prometheus", "Grafana"],
        "summary": ("Seasoned DevOps Engineer with 5 years of experience designing and "
                    "managing cloud-native infrastructure. Expert in IaC, container orchestration, "
                    "and SRE practices. Reduced deployment failures by 70% at previous role."),
        "jobs": [
            ("Flipkart", "Senior DevOps Engineer", "Jan 2022 – Present",
             ["Orchestrated Kubernetes clusters (200+ nodes) for Big Billion Day sales events",
              "Implemented Terraform modules for multi-region AWS infrastructure as code",
              "Built GitOps pipeline using ArgoCD reducing release cycle from 2 weeks to 1 day",
              "Set up Prometheus + Grafana observability stack with 150+ custom alerts"]),
            ("HCL Technologies", "DevOps Engineer", "Jun 2019 – Dec 2021",
             ["Managed Jenkins CI/CD for 40+ microservices with zero-downtime deployments",
              "Containerised 15 legacy apps using Docker, saving $120K/yr in infra costs"]),
        ],
        "certs": ["CKA – Certified Kubernetes Administrator",
                  "AWS DevOps Engineer Professional",
                  "HashiCorp Terraform Associate"],
        "format": "pdf",
    },
    {
        "name": "Ananya Krishnan",
        "email": "ananya.k@gmail.com",
        "phone": "+91 75019 66321",
        "location": "Chennai, Tamil Nadu",
        "role": "Product Manager",
        "experience": "7 years",
        "education": "MBA — IIM Calcutta, 2017 | B.Tech CSE — NIT Trichy, 2015",
        "skills": ["Product Strategy", "Roadmapping", "Agile", "Scrum", "JIRA", "Confluence",
                   "User Research", "SQL", "Mixpanel", "Amplitude", "Figma", "A/B Testing"],
        "summary": ("Strategic Product Manager with 7 years building B2B SaaS products. "
                    "Launched 4 major product lines, growing ARR from $2M to $18M. "
                    "Strong mix of technical background and customer empathy."),
        "jobs": [
            ("Freshworks", "Senior Product Manager", "Apr 2020 – Present",
             ["Owned Freshdesk's AI automation suite — grew MAU by 210% in 18 months",
              "Led cross-functional squad of 14 (eng, design, data, QA)",
              "Defined product OKRs and quarterly roadmaps aligned with company strategy",
              "Conducted 200+ customer interviews to prioritise feature backlog"]),
            ("Tata Consultancy Services", "Product Analyst", "Aug 2017 – Mar 2020",
             ["Analysed user behaviour data to inform product decisions for 3M-user platform",
              "Coordinated sprint planning and backlog grooming across 2 scrum teams"]),
        ],
        "certs": ["Certified Scrum Product Owner (CSPO)",
                  "Google Project Management Certificate"],
        "format": "docx",
    },
    {
        "name": "Vikram Singh",
        "email": "vikram.singh99@gmail.com",
        "phone": "+91 81200 34569",
        "location": "Noida, Uttar Pradesh",
        "role": "Android Developer",
        "experience": "3.5 years",
        "education": "B.Tech CSE — Amity University, 2020",
        "skills": ["Kotlin", "Java", "Android SDK", "Jetpack Compose", "MVVM", "Retrofit",
                   "Room DB", "Firebase", "Coroutines", "Hilt", "Material Design", "Git"],
        "summary": ("Android Developer with 3.5 years building high-quality native apps. "
                    "Shipped 6 apps to the Play Store with 4.5+ ratings. Passionate about "
                    "clean architecture, smooth animations, and offline-first design."),
        "jobs": [
            ("Paytm", "Android Developer", "Oct 2020 – Present",
             ["Maintained Paytm Money app (2M+ downloads) — improved crash rate from 1.4% to 0.2%",
              "Migrated XML layouts to Jetpack Compose, reducing UI code by 35%",
              "Implemented offline-first architecture using Room + WorkManager",
              "Built in-app investment calculator used by 800K+ users/month"]),
        ],
        "certs": ["Associate Android Developer – Google",
                  "Jetpack Compose Pathway – Android Developers"],
        "format": "pdf",
    },
    {
        "name": "Meera Nair",
        "email": "meera.nair@gmail.com",
        "phone": "+91 94400 12345",
        "location": "Kochi, Kerala",
        "role": "UI/UX Designer",
        "experience": "4 years",
        "education": "B.Des Interaction Design — NID Ahmedabad, 2020",
        "skills": ["Figma", "Adobe XD", "Sketch", "Prototyping", "User Research",
                   "Wireframing", "Design Systems", "Usability Testing", "HTML", "CSS",
                   "Zeplin", "Miro", "Accessibility (WCAG)"],
        "summary": ("UX Designer with a human-centred approach to solving complex design "
                    "problems. Designed interfaces for 3M+ users across fintech and edtech. "
                    "Strong believer in data-driven design decisions and inclusive design."),
        "jobs": [
            ("BYJU'S", "Senior UI/UX Designer", "Jun 2021 – Present",
             ["Led redesign of BYJU's student dashboard — increased session duration by 27%",
              "Built and maintained a 400+ component design system in Figma",
              "Conducted 80+ usability tests and synthesised findings into design improvements",
              "Collaborated daily with 6 product squads across 3 time zones"]),
            ("Mswipe", "UI Designer", "Jul 2020 – May 2021",
             ["Designed merchant-facing POS interface from wireframes to handoff",
              "Created interactive prototypes for stakeholder presentations"]),
        ],
        "certs": ["Google UX Design Certificate",
                  "Interaction Design Foundation – UX Research"],
        "format": "docx",
    },
    {
        "name": "Siddharth Rao",
        "email": "sid.rao@gmail.com",
        "phone": "+91 73000 98765",
        "location": "Bengaluru, Karnataka",
        "role": "Backend Engineer (Fresher)",
        "experience": "0 years (recent graduate)",
        "education": "B.E. Computer Science — RV College of Engineering, 2025",
        "skills": ["Python", "Java", "Spring Boot", "SQL", "MongoDB", "REST APIs",
                   "Git", "Linux", "HTML", "CSS", "DSA"],
        "summary": ("Recent B.E. graduate with strong fundamentals in data structures and "
                    "algorithms, backend development, and database design. Completed 2 internships "
                    "and 4 personal projects. Eager to join a fast-paced engineering team."),
        "jobs": [
            ("Accenture (Intern)", "Software Engineering Intern", "Jan 2025 – Apr 2025",
             ["Built REST APIs using Spring Boot for an internal HR management tool",
              "Wrote JUnit tests achieving 82% code coverage",
              "Participated in daily standups and sprint retrospectives"]),
            ("Startup (Internship)", "Backend Intern", "Jun 2024 – Sep 2024",
             ["Designed MongoDB schema for a logistics tracking app",
              "Implemented JWT authentication and role-based access control"]),
        ],
        "certs": ["HackerRank Python (Gold Badge)",
                  "CS50x — Harvard University (edX)"],
        "format": "pdf",
    },
    {
        "name": "Divya Patel",
        "email": "divya.patel@gmail.com",
        "phone": "+91 87654 32100",
        "location": "Ahmedabad, Gujarat",
        "role": "Cloud Solutions Architect",
        "experience": "9 years",
        "education": "M.Tech CSE — IIT Delhi, 2016 | B.E. CSE — NIRMA University, 2014",
        "skills": ["AWS", "Azure", "GCP", "Terraform", "CloudFormation", "Kubernetes",
                   "Python", "Microservices", "Serverless", "Cost Optimisation",
                   "Security (IAM, VPC)", "CDN", "RDS", "DynamoDB"],
        "summary": ("Cloud Architect with 9 years designing enterprise-grade, secure, and "
                    "cost-efficient cloud infrastructure on AWS and Azure. Led $15M digital "
                    "transformation for a major bank. 12 active cloud certifications."),
        "jobs": [
            ("HDFC Bank (Consultant)", "Principal Cloud Architect", "Mar 2020 – Present",
             ["Architected AWS-native banking platform processing ₹40,000 Cr daily",
              "Reduced cloud spend by 34% ($2.1M/yr) through Reserved Instances and rightsizing",
              "Designed zero-trust network architecture compliant with RBI guidelines",
              "Led team of 8 cloud engineers and 3 security specialists"]),
            ("Wipro", "Cloud Solutions Architect", "Apr 2016 – Feb 2020",
             ["Delivered cloud migration for 40+ enterprise clients across BFSI and retail",
              "Defined cloud CoE standards adopted across Wipro's 15,000 cloud practice"]),
        ],
        "certs": ["AWS Solutions Architect Professional",
                  "AWS Security Specialty",
                  "Azure Solutions Architect Expert",
                  "GCP Professional Cloud Architect"],
        "format": "docx",
    },
]


# ── PDF generator ─────────────────────────────────────────────────────────────
def make_pdf(candidate: dict, out_path: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    HRFlowable, ListFlowable, ListItem)

    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=1.8*cm, bottomMargin=1.8*cm)

    styles = getSampleStyleSheet()
    DARK   = colors.HexColor("#1a1a2e")
    BLUE   = colors.HexColor("#1f77b4")
    GREY   = colors.HexColor("#555555")
    LGREY  = colors.HexColor("#888888")

    name_style    = ParagraphStyle("Name",    fontSize=22, fontName="Helvetica-Bold",
                                   textColor=DARK, spaceAfter=4)
    role_style    = ParagraphStyle("Role",    fontSize=12, fontName="Helvetica",
                                   textColor=BLUE, spaceAfter=2)
    contact_style = ParagraphStyle("Contact", fontSize=9,  fontName="Helvetica",
                                   textColor=LGREY, spaceAfter=6)
    section_style = ParagraphStyle("Section", fontSize=11, fontName="Helvetica-Bold",
                                   textColor=DARK, spaceBefore=10, spaceAfter=4)
    body_style    = ParagraphStyle("Body",    fontSize=9.5, fontName="Helvetica",
                                   textColor=GREY, leading=14, spaceAfter=4)
    job_co_style  = ParagraphStyle("JCo",    fontSize=10, fontName="Helvetica-Bold",
                                   textColor=DARK, spaceBefore=6, spaceAfter=1)
    job_dt_style  = ParagraphStyle("JDt",    fontSize=9,  fontName="Helvetica-Oblique",
                                   textColor=LGREY, spaceAfter=3)
    bullet_style  = ParagraphStyle("Bul",    fontSize=9.5, fontName="Helvetica",
                                   textColor=GREY, leading=14)

    story = []

    # Header
    story.append(Paragraph(candidate["name"], name_style))
    story.append(Paragraph(candidate["role"], role_style))
    story.append(Paragraph(
        f'{candidate["email"]}  •  {candidate["phone"]}  •  {candidate["location"]}',
        contact_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dddddd")))
    story.append(Spacer(1, 6))

    # Summary
    story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
    story.append(Paragraph(candidate["summary"], body_style))

    # Experience
    story.append(Paragraph("WORK EXPERIENCE", section_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#eeeeee")))
    for company, title, dates, bullets in candidate["jobs"]:
        story.append(Paragraph(f"{company} — {title}", job_co_style))
        story.append(Paragraph(dates, job_dt_style))
        items = [ListItem(Paragraph(b, bullet_style), leftIndent=10, bulletColor=BLUE)
                 for b in bullets]
        story.append(ListFlowable(items, bulletType="bullet", leftIndent=12,
                                  bulletFontSize=8, start="•"))
        story.append(Spacer(1, 4))

    # Education
    story.append(Paragraph("EDUCATION", section_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#eeeeee")))
    story.append(Paragraph(candidate["education"], body_style))

    # Skills
    story.append(Paragraph("SKILLS", section_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#eeeeee")))
    story.append(Paragraph("  •  ".join(candidate["skills"]), body_style))

    # Certifications
    if candidate.get("certs"):
        story.append(Paragraph("CERTIFICATIONS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#eeeeee")))
        for cert in candidate["certs"]:
            story.append(Paragraph(f"✓  {cert}", body_style))

    doc.build(story)
    print(f"  [PDF] {out_path.name}")


# ── DOCX generator ────────────────────────────────────────────────────────────
def make_docx(candidate: dict, out_path: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margins
    for section in doc.sections:
        section.left_margin  = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin   = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    def add_para(text="", bold=False, italic=False, size=11,
                 color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_hr():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "DDDDDD")
        pBdr.append(bottom)
        pPr.append(pBdr)

    DARK  = (26,  26,  46)
    BLUE  = (31,  119, 180)
    GREY  = (85,  85,  85)
    LGREY = (136, 136, 136)

    # Header
    add_para(candidate["name"], bold=True, size=22, color=DARK, space_after=2)
    add_para(candidate["role"], size=13, color=BLUE, space_after=2)
    add_para(f'{candidate["email"]}  •  {candidate["phone"]}  •  {candidate["location"]}',
             size=9, color=LGREY, space_after=6)
    add_hr()

    # Summary
    add_para("PROFESSIONAL SUMMARY", bold=True, size=11, color=DARK, space_after=3)
    add_para(candidate["summary"], size=9, color=GREY, space_after=8)

    # Experience
    add_para("WORK EXPERIENCE", bold=True, size=11, color=DARK, space_after=3)
    add_hr()
    for company, title, dates, bullets in candidate["jobs"]:
        add_para(f"{company} — {title}", bold=True, size=10, color=DARK, space_after=1)
        add_para(dates, italic=True, size=9, color=LGREY, space_after=3)
        for b in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(b)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(*GREY)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Education
    add_para("EDUCATION", bold=True, size=11, color=DARK, space_after=3)
    add_hr()
    add_para(candidate["education"], size=9, color=GREY, space_after=8)

    # Skills
    add_para("SKILLS", bold=True, size=11, color=DARK, space_after=3)
    add_hr()
    add_para("  •  ".join(candidate["skills"]), size=9, color=GREY, space_after=8)

    # Certifications
    if candidate.get("certs"):
        add_para("CERTIFICATIONS", bold=True, size=11, color=DARK, space_after=3)
        add_hr()
        for cert in candidate["certs"]:
            add_para(f"✓  {cert}", size=9, color=GREY, space_after=3)

    doc.save(str(out_path))
    print(f"  [DOCX] {out_path.name}")


# ── Generate all resumes ──────────────────────────────────────────────────────
generated = []
for c in CANDIDATES:
    safe_name = c["name"].lower().replace(" ", "_")
    if c["format"] == "pdf":
        path = OUT_DIR / f"{safe_name}_resume.pdf"
        make_pdf(c, path)
    else:
        path = OUT_DIR / f"{safe_name}_resume.docx"
        make_docx(c, path)
    generated.append(path)

# ── Create ZIP ────────────────────────────────────────────────────────────────
zip_path = OUT_DIR / "test_resumes_10.zip"
with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
    for p in generated:
        zf.write(p, p.name)

print(f"\nDone! {len(generated)} resumes + ZIP created in: {OUT_DIR}")
print(f"  ZIP: {zip_path}")
for p in generated:
    print(f"  {p.name}")
