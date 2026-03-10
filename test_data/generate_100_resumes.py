import os
import random
import zipfile
from docx import Document

# List of tech skills to randomly assign
SKILLS_DB = [
    "Python", "Java", "C++", "JavaScript", "React", "Node.js", "AWS", "Docker",
    "Kubernetes", "SQL", "MongoDB", "PostgreSQL", "Machine Learning", "Data Science",
    "TensorFlow", "PyTorch", "DevOps", "CI/CD", "Git", "Agile", "Scrum",
    "UI/UX", "Figma", "HTML", "CSS", "TypeScript", "Angular", "Vue.js",
    "Spring Boot", "Django", "Flask", "FastAPI", "Ruby on Rails", "C#", ".NET"
]

# List of names
FIRST_NAMES = ["James", "Mary", "John", "Patricia", "Robert", "Jennifer", "Michael", "Linda", "William", "Elizabeth", "David", "Barbara", "Richard", "Susan", "Joseph", "Jessica", "Thomas", "Sarah", "Charles", "Karen"]
LAST_NAMES = ["Smith", "Johnson", "Williams", "Brown", "Jones", "Garcia", "Miller", "Davis", "Rodriguez", "Martinez", "Hernandez", "Lopez", "Gonzalez", "Wilson", "Anderson", "Thomas", "Taylor", "Moore", "Jackson", "Martin"]

def generate_resume(file_path, name, skills, is_anomaly=False):
    doc = Document()
    doc.add_heading(f'{name} - Resume', 0)
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Experienced software engineer with a track record of building scalable web applications and data pipelines.')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph(', '.join(skills))
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Senior Software Engineer - Tech Corp (2018-Present)\n- Developed amazing features using modern web technologies.')
    doc.add_paragraph('Software Engineer - Startup Inc (2015-2018)\n- Maintained critical infrastructure and improved performance.')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.S. in Computer Science - University of Technology')
    
    # If generating an anomaly, simulate keyword stuffing by repeating skills a massive number of times
    if is_anomaly:
        doc.add_heading('Hidden Skills section for SEO', level=1)
        stuffing = " ".join(skills * 100) # Repeat skills 100 times to create an abnormally long document
        doc.add_paragraph(stuffing)
        
    doc.save(file_path)
    return file_path

def main():
    output_dir = "mock_100_resumes"
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"Generating 100 resumes...")
    
    generated_files = []
    
    for i in range(1, 101):
        name = f"{random.choice(FIRST_NAMES)} {random.choice(LAST_NAMES)}"
        
        # Pick 3 to 10 random skills for each person
        num_skills = random.randint(3, 10)
        skills = random.sample(SKILLS_DB, num_skills)
        
        # 5% chance to create an anomaly
        is_anomaly = random.random() < 0.05
        
        # Create a safe filename
        safe_name = name.lower().replace(" ", "_")
        # Add index to avoid naming collisions
        filename = f"{safe_name}_{i}.docx"
        filepath = os.path.join(output_dir, filename)
        
        generate_resume(filepath, name, skills, is_anomaly)
        generated_files.append(filepath)
        
    zip_path = "100_resumes.zip"
    print(f"Zipping {len(generated_files)} resumes into {zip_path}...")
    
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for file in generated_files:
            # Add to zip without the outer physical directory structure
            zipf.write(file, os.path.basename(file))
            
    print(f"Done! Created {zip_path}")
    
    # Clean up the unzipped files to save space
    for file in generated_files:
        os.remove(file)
    os.rmdir(output_dir)

if __name__ == "__main__":
    main()
